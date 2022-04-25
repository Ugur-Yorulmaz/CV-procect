import docx

from docx import Document
from docx.shared import Inches
import pyttsx3

# Speaking example

def speak(text):
    pyttsx3.speak(text)

speak('Ugur Yorulmaz')

#CV Hazırlama Programı

document=Document()
# Profile picture
document.add_picture ('resim.jpg',width=Inches(2.0))

# Name phone and e-mail
name=input('Adınız nedir ? ')
phone=input('Telefonunuz nedir ? ')
email=input('emailiniz nedir ? ')

document.add_paragraph (name + '|'+phone +'|' + email)

#About me
document.add_heading ('About Me')
about_me = input('Kendi hakkında birşeyler yaz   ')
document.add_paragraph (about_me)

# Work Experience
document.add_heading ('Work Experience')

print ('Work Experiences')
# Entering more experiences
while True:


    p=document.add_paragraph()

    company=input ('Company ')
    from_date = input('From Date ')
    to_date = input ('To date ')

    p.add_run(company + ' ').bold = True
    p.add_run (from_date + '-' + to_date +'\n').italic = True

    experience_details = input (' Enter your experience details in' + company)
    p.add_run (experience_details)

    i=input('Do you have any other work experience ?')

    if i.lower()=='yes':
        continue
    else:
        break
# Entering skills
document.add_heading('Skills')

# Entering more skills
while True:
    skill = input('Enter your skill')
    p=document.add_paragraph (skill)
    k=input('Do you like to enter more skill ? ')


    if k.lower()=='yes':
        continue
    else:
        break

document.add_paragraph

# Entering footer
section=document.section(0)
footer=section.footer
p=footer.paragraph(0)
p_text='This code has written by Ugur using Amigoscode'

# Saving document
document.save('cv.docx')


